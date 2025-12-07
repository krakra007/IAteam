import streamlit as st
import os
import signal 
from dotenv import load_dotenv
import google.generativeai as genai
from openai import OpenAI
from anthropic import Anthropic 
import concurrent.futures 
import PyPDF2 
import docx 
from docx import Document 
from io import BytesIO
import time
import uuid
from st_copy_to_clipboard import st_copy_to_clipboard

# 1. Chargement des clés API
load_dotenv()

# ==========================================
# CONFIGURATION
# ==========================================
try:
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
    model_gemini = genai.GenerativeModel('gemini-2.5-flash')
except Exception as e:
    st.error(f"Erreur Config Gemini : {e}")

client_openai = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

client_perplexity = OpenAI(
    api_key=os.getenv("PERPLEXITY_API_KEY"),
    base_url="https://api.perplexity.ai"
)

try:
    client_claude = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
except Exception as e:
    st.error(f"Erreur Config Claude : {e}")

# ==========================================
# GESTION D'ÉTAT
# ==========================================
if "ai_states" not in st.session_state:
    st.session_state.ai_states = {
        "gemini": True,
        "chatgpt": True,
        "perplexity": True,
        "claude": True
    }

def toggle_ai(ai_name):
    st.session_state.ai_states[ai_name] = not st.session_state.ai_states[ai_name]

# ==========================================
# UTILITAIRES (CSS "GRAND CONFORT" 18px)
# ==========================================
def inject_custom_css():
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap');
        
        html, body, [class*="css"] { font-family: 'Google Sans', sans-serif !important; }
        
        /* === MODIFICATION TAILLE POLICE (18px) === */
        p, .stMarkdown, .stText, li, div { 
            font-size: 18px !important; 
            line-height: 1.7 !important; /* Un peu plus d'espace entre les lignes */
            font-weight: 400 !important; 
        }
        
        h1, h2, h3 { font-family: 'Google Sans', sans-serif !important; font-weight: 600 !important; }
        .main .block-container { max_width: 95% !important; padding-bottom: 8rem !important; }
        
        /* INPUT DU BAS (Zone de saisie) */
        [data-testid="stChatInput"] { border: none !important; background: transparent !important; }
        [data-testid="stChatInput"] > div { background-color: transparent !important; border-color: transparent !important; box-shadow: none !important; }
        [data-testid="stChatInput"] textarea {
            height: 150px !important; 
            min-height: 150px !important; 
            font-size: 18px !important; /* <-- 18px ICI AUSSI */
            padding: 15px !important; 
            border: 1px solid rgba(128, 128, 128, 0.4) !important;
            border-radius: 12px !important; 
            background-color: transparent !important;
            font-family: 'Google Sans', sans-serif !important;
        }
        [data-testid="stChatInput"] textarea:focus { border: 1px solid #777 !important; box-shadow: none !important; }

        /* BOUTONS SIDEBAR */
        [data-testid="stSidebar"] button { font-size: 13px !important; padding: 0.25rem 0.5rem !important; font-weight: 600 !important; }
        [data-testid="stSidebar"] button[kind="primary"] { background-color: #2e7d32 !important; border-color: #2e7d32 !important; color: white !important; }
        [data-testid="stSidebar"] button[kind="secondary"] { border: 1px solid #444 !important; color: #888 !important; background-color: transparent !important; }
    </style>
    """, unsafe_allow_html=True)

def read_uploaded_file(uploaded_file):
    if uploaded_file is None: return ""
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages: text += page.extract_text() + "\n"
            return text
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(uploaded_file)
            return '\n'.join([para.text for para in doc.paragraphs])
        else: return uploaded_file.getvalue().decode("utf-8")
    except Exception as e: return f"Erreur lecture : {e}"

def generate_docx(markdown_text):
    doc = Document()
    doc.add_heading('Rapport IA Expert', 0)
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('- '): doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        else: doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def stop_program():
    os.kill(os.getpid(), signal.SIGTERM)

def reset_conversation():
    st.session_state.messages = []
    st.session_state.chat_id = str(uuid.uuid4())
    st.rerun()

def get_conversation_history():
    history_text = ""
    if "messages" in st.session_state:
        recent_messages = st.session_state.messages[-6:] 
        for msg in recent_messages:
            role = "UTILISATEUR" if msg["role"] == "user" else "ASSISTANT (Réponse précédente)"
            content = msg["content"][:2000] 
            history_text += f"{role}: {content}\n\n"
    return history_text

# ==========================================
# RECHERCHE (AVEC MÉMOIRE PARTAGÉE)
# ==========================================

def ask_gemini_generalist(current_prompt, history):
    try:
        full_prompt = f"HISTORIQUE DE CONVERSATION:\n{history}\n\nNOUVELLE QUESTION: {current_prompt}\nCONSIGNE: Réponds en tenant compte de l'historique si pertinent. Markdown & LaTeX ($...$)."
        response = model_gemini.generate_content(full_prompt)
        return response.text
    except Exception as e: return f"Erreur Gemini : {e}"

def ask_chatgpt_generalist(current_prompt, history):
    try:
        full_content = f"CONTEXTE PRÉCÉDENT:\n{history}\n\nQUESTION ACTUELLE:\n{current_prompt}"
        sys_msg = "Expert Universel. Markdown & LaTeX."
        response = client_openai.chat.completions.create(
            model="gpt-4o", 
            messages=[{"role": "system", "content": sys_msg}, {"role": "user", "content": full_content}],
            max_tokens=3000
        )
        return response.choices[0].message.content
    except Exception as e: return f"Erreur ChatGPT : {e}"

def ask_perplexity_generalist(current_prompt, history):
    try:
        full_content = f"CONTEXTE PRÉCÉDENT:\n{history}\n\nQUESTION ACTUELLE:\n{current_prompt}"
        sys_msg = "Expert Web. URLs à la fin."
        response = client_perplexity.chat.completions.create(
            model="sonar-pro", 
            messages=[{"role": "system", "content": sys_msg}, {"role": "user", "content": full_content}]
        )
        return response.choices[0].message.content
    except Exception as e: return f"Erreur Perplexity : {e}"

def ask_claude_generalist(current_prompt, history):
    if client_claude is None: return "⚠️ Claude désactivé"
    full_content = f"HISTORIQUE:\n{history}\n\nQUESTION:\n{current_prompt}"
    models_to_try = ["claude-sonnet-4-5-20250929", "claude-3-5-sonnet-20241022", "claude-3-5-sonnet-20240620", "claude-3-opus-20240229"]
    for model_id in models_to_try:
        try:
            message = client_claude.messages.create(
                model=model_id, 
                max_tokens=3000,
                system="Expert Analytique. Markdown & LaTeX.",
                messages=[{"role": "user", "content": full_content}]
            )
            return message.content[0].text
        except Exception as e:
            if "not_found" in str(e): continue
            else: return f"Erreur Claude : {e}"
    return "Erreur Claude : Aucun modèle trouvé."

# ==========================================
# COMPILATION
# ==========================================

def compile_quad_fusion(original_prompt, responses, file_content, history):
    
    bloc_fichier = f"\n=== DOCUMENT LOCAL ===\n{file_content[:40000]}\n" if file_content else ""
    
    responses_text = ""
    for name, content in responses.items():
        responses_text += f"\n=== AVIS DE {name.upper()} ===\n{content}\n"

    compilation_prompt = f"""
    HISTORIQUE GLOBAL : {history}
    QUESTION ACTUELLE : "{original_prompt}"
    
    TU ES UN ARCHITECTE DE L'INFORMATION.
    MISSION : Fusionner les réponses des experts ci-dessous en un document unique.
    
    RÈGLES :
    1. **MARKDOWN** pour le texte.
    2. **LATEX** pour les maths ($...$).
    3. **CONTINUITÉ** : Si la question fait référence au passé, utilise l'historique.
    
    {responses_text}
    
    {bloc_fichier}
    
    Génère le document final.
    """
    
    try:
        response = model_gemini.generate_content(compilation_prompt)
        return response.text
    except Exception as e: return f"Erreur Compilation : {e}"

# ==========================================
# INTERFACE
# ==========================================

st.set_page_config(page_title="Super IA Universelle", page_icon="🧠", layout="wide")
inject_custom_css() 

st.title("🧠 Super IA Universelle")

if "chat_id" not in st.session_state:
    st.session_state.chat_id = str(uuid.uuid4())

with st.sidebar:
    st.header("🎛️ Moteurs")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("GEMINI", type="primary" if st.session_state.ai_states["gemini"] else "secondary", use_container_width=True):
            toggle_ai("gemini")
            st.rerun()
        if st.button("CLAUDE", type="primary" if st.session_state.ai_states["claude"] else "secondary", use_container_width=True):
            toggle_ai("claude")
            st.rerun()
    with col2:
        if st.button("CHATGPT", type="primary" if st.session_state.ai_states["chatgpt"] else "secondary", use_container_width=True):
            toggle_ai("chatgpt")
            st.rerun()
        if st.button("PERPLEXITY", type="primary" if st.session_state.ai_states["perplexity"] else "secondary", use_container_width=True):
            toggle_ai("perplexity")
            st.rerun()

    st.markdown("---")
    st.header("📂 Documents")
    uploaded_file = st.file_uploader("Ajouter contexte", type=["pdf", "docx", "txt"])
    user_file_content = read_uploaded_file(uploaded_file)
    if user_file_content: st.success("Fichier chargé.")
    
    st.markdown("---")
    st.header("⚙️ Contrôles")
    if st.button("🧹 Reset"):
        reset_conversation()
    if st.button("🛑 Stop"):
        st.warning("Arrêt...")
        time.sleep(1)
        stop_program()

if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

if prompt := st.chat_input("Posez votre question...", key=st.session_state.chat_id):
    
    active_ais = [k for k, v in st.session_state.ai_states.items() if v]
    if not active_ais:
        st.error("⚠️ Activez au moins une IA !")
    else:
        history_str = get_conversation_history()
        
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            status = st.status("🚀 Recherche en cours (avec contexte)...", expanded=True)
            responses = {}
            
            with concurrent.futures.ThreadPoolExecutor() as executor:
                futures = {}
                if st.session_state.ai_states["gemini"]: 
                    futures["Gemini"] = executor.submit(ask_gemini_generalist, prompt, history_str)
                if st.session_state.ai_states["chatgpt"]: 
                    futures["ChatGPT"] = executor.submit(ask_chatgpt_generalist, prompt, history_str)
                if st.session_state.ai_states["perplexity"]: 
                    futures["Perplexity"] = executor.submit(ask_perplexity_generalist, prompt, history_str)
                if st.session_state.ai_states["claude"]: 
                    futures["Claude"] = executor.submit(ask_claude_generalist, prompt, history_str)
                
                for name, future in futures.items():
                    result = future.result()
                    responses[name] = result
                    if "Erreur" in result: status.write(f"⚠️ {name}")
                    else: status.write(f"✅ {name}")
                
                status.update(label="🧬 Fusion...", state="running")
                final_doc = compile_quad_fusion(prompt, responses, user_file_content, history_str)
                status.update(label="Terminé !", state="complete", expanded=False)

            st.markdown(final_doc)
            
            col1, col2, col3 = st.columns([1, 1, 8])
            with col1: st_copy_to_clipboard(final_doc, before_copy_label="📋", after_copy_label="✅")
            with col2: st.download_button("📝", data=final_doc, file_name="rapport.md", mime="text/markdown")
            with col3: 
                docx_file = generate_docx(final_doc)
                st.download_button("📘", data=docx_file, file_name="rapport.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            st.markdown("---")
            with st.expander("🧐 Réponses individuelles"):
                if responses:
                    tabs = st.tabs(list(responses.keys()))
                    for i, (name, content) in enumerate(responses.items()):
                        with tabs[i]: st.markdown(content)
            
        st.session_state.messages.append({"role": "assistant", "content": final_doc})
